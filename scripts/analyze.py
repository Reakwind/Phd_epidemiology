import argparse
import os
from typing import Dict, List

import pandas as pd
import matplotlib.pyplot as plt
from docx import Document


def load_data(file_path: str) -> pd.DataFrame:
    """Load a CSV file into a DataFrame."""
    return pd.read_csv(file_path)


def clean_data(df: pd.DataFrame) -> pd.DataFrame:
    """\
    Basic cleaning tailored for medical research data.

    - Drops duplicate rows.
    - Attempts to convert columns to numeric when possible.
    - Fills missing values: numeric columns get the median, others get the mode.
    """

    df = df.drop_duplicates()

    # Try to convert columns that look numeric; leave categorical data intact
    for col in df.columns:
        converted = pd.to_numeric(df[col], errors="coerce")
        if converted.notna().sum() > 0:
            df[col] = converted

    for col in df.columns:
        if pd.api.types.is_numeric_dtype(df[col]):
            df[col] = df[col].fillna(df[col].median())
        else:
            mode = df[col].mode()
            df[col] = df[col].fillna(mode.iloc[0] if not mode.empty else "missing")

    return df


def analyze_data(df: pd.DataFrame) -> Dict[str, pd.DataFrame]:
    """\
    Generate descriptive statistics for medical datasets.

    Returns a dictionary with numeric summaries, categorical counts, and a
    correlation matrix for numeric variables.
    """

    summaries: Dict[str, pd.DataFrame] = {}

    numeric = df.select_dtypes(include="number")
    if not numeric.empty:
        summaries["numeric"] = numeric.describe().T
        summaries["correlation"] = numeric.corr()

    categorical_columns = df.select_dtypes(exclude="number").columns
    categorical_summary = {}
    for col in categorical_columns:
        categorical_summary[col] = df[col].value_counts().to_frame(name="count")
    if categorical_summary:
        summaries["categorical"] = pd.concat(categorical_summary, axis=0)

    return summaries


def generate_visualizations(df: pd.DataFrame, output_dir: str) -> List[str]:
    """Create plots for numeric and categorical columns."""
    figure_paths: List[str] = []

    numeric = df.select_dtypes(include="number")
    if not numeric.empty:
        path = os.path.join(output_dir, "numeric.png")
        numeric.hist(figsize=(10, 8))
        plt.tight_layout()
        plt.savefig(path)
        plt.close()
        figure_paths.append(path)

    categorical = df.select_dtypes(exclude="number")
    for col in categorical.columns:
        path = os.path.join(output_dir, f"{col}_bar.png")
        df[col].value_counts().plot(kind="bar")
        plt.title(col)
        plt.tight_layout()
        plt.savefig(path)
        plt.close()
        figure_paths.append(path)

    return figure_paths


def _add_dataframe_table(document: Document, df: pd.DataFrame) -> None:
    """Add a pandas DataFrame to a Word document as a table."""
    table = document.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def export_to_docx(
    summaries: Dict[str, pd.DataFrame], figure_paths: List[str], output_docx: str
) -> None:
    """Save summary statistics and figures to a docx file."""
    document = Document()
    document.add_heading("Analysis Summary", level=1)

    numeric = summaries.get("numeric")
    if numeric is not None:
        document.add_heading("Numeric Variables", level=2)
        _add_dataframe_table(document, numeric.reset_index())

    categorical = summaries.get("categorical")
    if categorical is not None:
        document.add_heading("Categorical Variables", level=2)
        _add_dataframe_table(document, categorical.reset_index())

    correlation = summaries.get("correlation")
    if correlation is not None:
        document.add_heading("Correlation Matrix", level=2)
        _add_dataframe_table(document, correlation.reset_index())

    for path in figure_paths:
        if os.path.exists(path):
            document.add_picture(path)

    document.save(output_docx)


def main(csv_file: str, output_dir: str) -> None:
    """Run the analysis pipeline on `csv_file` and save outputs."""
    if not os.path.isfile(csv_file):
        raise FileNotFoundError(csv_file)

    os.makedirs(output_dir, exist_ok=True)

    df = load_data(csv_file)
    df = clean_data(df)
    summaries = analyze_data(df)

    figure_paths = generate_visualizations(df, output_dir)

    docx_path = os.path.join(output_dir, "report.docx")
    export_to_docx(summaries, figure_paths, docx_path)

    print(f"Report saved to {docx_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Analyze tabular medical research data and generate a docx report"
    )
    parser.add_argument("csv_file", help="Path to input CSV file")
    parser.add_argument("output_dir", help="Directory to save outputs")
    args = parser.parse_args()
    main(args.csv_file, args.output_dir)
